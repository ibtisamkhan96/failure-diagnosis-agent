"""Export the agent's Markdown diagnosis to a formatted Word (.docx) failure-analysis report.

Parses the headings, bullets, numbered lists, and bold spans the agent produces, and wraps them
in a titled report with a clear disclaimer that this is a preliminary screening report, not a
completed physical examination.
"""

import datetime
import re


def _add_inline(paragraph, text):
    """Add text to a paragraph, rendering **bold** spans."""
    for part in re.split(r"(\*\*.+?\*\*)", text):
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            paragraph.add_run(part[2:-2]).bold = True
        elif part:
            paragraph.add_run(part)


def markdown_to_docx(md, out_path, title="Failure Analysis Report (Preliminary)",
                     author="Ibtisam Ahmed Khan", used_photo=False):
    from docx import Document  # imported here so the rest of the project needs no docx install

    doc = Document()
    doc.add_heading(title, level=0)

    meta = doc.add_paragraph()
    meta.add_run(f"Prepared by: {author}").italic = True
    meta.add_run(f"     Date: {datetime.date.today().isoformat()}").italic = True

    src = "the written description" + (" and the photograph" if used_photo else "") + " provided"
    disc = doc.add_paragraph()
    disc.add_run(
        "This is a preliminary screening report produced by an AI failure-diagnosis agent from "
        f"{src}. It is a structured reasoning aid, not a completed physical examination or "
        "laboratory failure analysis. The recommended tests have not been performed. Conclusions "
        "should be confirmed by hands-on examination and the cited standard tests before any action."
    ).italic = True

    doc.add_paragraph()  # spacer

    for raw in md.split("\n"):
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif re.match(r"^\s*[-*]\s+", line):
            _add_inline(doc.add_paragraph(style="List Bullet"), re.sub(r"^\s*[-*]\s+", "", line))
        elif re.match(r"^\s*\d+\.\s+", line):
            _add_inline(doc.add_paragraph(style="List Number"), re.sub(r"^\s*\d+\.\s+", "", line))
        else:
            _add_inline(doc.add_paragraph(), line)

    doc.save(out_path)
    return out_path
